import pandas as pd
import streamlit as st
from docx import Document
import os
import re
import sys
import zipfile
import io
import tempfile
from typing import Dict, List
from pathlib import Path
import datetime

# 设置页面标题
st.set_page_config(page_title="文档批量生成工具", layout="wide")
st.title("商标文档批量生成工具")

# 初始化session状态
if 'processing_stage' not in st.session_state:
    st.session_state.processing_stage = 0  # 0: 未开始, 1: 处理完成
if 'output_dir' not in st.session_state:
    st.session_state.output_dir = ""
if 'generated_files' not in st.session_state:
    st.session_state.generated_files = []
if 'filename_template' not in st.session_state:
    st.session_state.filename_template = ""

# 占位符处理器
class PlaceholderHandler:
    """占位符处理器，支持多种占位符格式"""
    PLACEHOLDER_PATTERNS = [
        r'\{\{\s*(.*?)\s*\}\}',  # {{key}}
        r'\$\{\s*(.*?)\s*\}',    # ${key}
        r'\{\s*(.*?)\s*\}',      # {key}
        r'\[\[\s*(.*?)\s*\]\]'   # [[key]]
    ]

    @classmethod
    def find_placeholders(cls, text: str) -> List[str]:
        """查找文本中的所有占位符"""
        placeholders = []
        for pattern in cls.PLACEHOLDER_PATTERNS:
            matches = re.findall(pattern, text)
            placeholders.extend(matches)
        return list(set(placeholders))

    @classmethod
    def replace_placeholder(cls, text: str, placeholder: str, value: str) -> str:
        """替换特定格式的占位符"""
        for pattern in cls.PLACEHOLDER_PATTERNS:
            wrapped_placeholder = pattern.replace(r'(.*?)', re.escape(placeholder))
            if re.search(wrapped_placeholder, text):
                text = re.sub(wrapped_placeholder, str(value), text)
        return text

def replace_text_in_paragraph(paragraph, replacements: Dict[str, str]):
    """替换段落中的占位符，保留原有格式"""
    all_placeholders = []
    
    # 检查段落中是否包含任何占位符
    for key in replacements:
        for pattern in PlaceholderHandler.PLACEHOLDER_PATTERNS:
            wrapped_key = pattern.replace(r'(.*?)', re.escape(key))
            if re.search(wrapped_key, paragraph.text):
                all_placeholders.append(key)
                break
    
    if not all_placeholders:
        return
    
    # 获取段落的完整文本（合并所有run）
    full_text = ''.join([run.text for run in paragraph.runs])
    
    # 替换所有占位符
    for placeholder in all_placeholders:
        value = replacements.get(placeholder, '')
        for pattern in PlaceholderHandler.PLACEHOLDER_PATTERNS:
            wrapped_placeholder = pattern.replace(r'(.*?)', re.escape(placeholder))
            full_text = re.sub(wrapped_placeholder, str(value), full_text)
    
    # 清空原有runs并添加新文本
    for run in paragraph.runs:
        run.text = ""
    
    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)

def process_document(template_path: str, output_path: str, replacements: Dict[str, str]) -> bool:
    """处理整个Word文档的替换"""
    try:
        doc = Document(template_path)
        
        # 处理正文段落
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
        
        # 处理页眉
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            
            # 处理页脚
            for paragraph in section.footer.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
        
        doc.save(output_path)
        return True
    except Exception as e:
        st.error(f"处理文档时发生错误: {str(e)}")
        return False

def generate_output_filename(row: Dict[str, str], filename_template: str) -> str:
    """使用模板生成输出文件名"""
    filename = filename_template
    
    # 替换模板中的占位符
    for key, value in row.items():
        for pattern in PlaceholderHandler.PLACEHOLDER_PATTERNS:
            wrapped_key = pattern.replace(r'(.*?)', re.escape(key))
            if re.search(wrapped_key, filename):
                filename = re.sub(wrapped_key, str(value), filename)
    
    # 添加日期时间戳避免重复
    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"{filename}_{timestamp}"
    
    # 移除非法字符
    filename = re.sub(r'[\\/*?:"<>|]', "", filename).strip()
    return f"{filename[:100]}.docx"  # 限制文件名长度

# 文件上传区域
st.header("1. 上传文件")

col1, col2 = st.columns(2)

with col1:
    excel_file = st.file_uploader("上传Excel数据文件", type=["xlsx", "xls"])
with col2:
    template_file = st.file_uploader("上传Word模板文件", type=["docx"])

# 文件名模板区域
st.header("2. 文件名模板")
filename_template = st.text_input(
    "文件名模板（使用与Word模板相同的占位符格式）",
    value="请款单（{{申请人}}-{{合计}}-集佳-{{日期}}）",
    help="示例：请款单（{{申请人}}-{{合计}}-集佳-{{日期}}）"
)

# 处理按钮
if st.button("开始生成文档") and excel_file and template_file:
    # 创建临时目录
    temp_dir = tempfile.mkdtemp()
    st.session_state.output_dir = os.path.join(temp_dir, "生成文档")
    os.makedirs(st.session_state.output_dir, exist_ok=True)
    
    # 保存上传的文件
    excel_path = os.path.join(temp_dir, excel_file.name)
    with open(excel_path, "wb") as f:
        f.write(excel_file.getbuffer())
    
    template_path = os.path.join(temp_dir, template_file.name)
    with open(template_path, "wb") as f:
        f.write(template_file.getbuffer())
    
    # 处理Excel数据
    try:
        # 读取Excel
        df = pd.read_excel(excel_path).astype(str)
        
        # 生成文档
        progress_bar = st.progress(0)
        status_text = st.empty()
        generated_files = []
        
        total_rows = len(df)
        success_count = 0
        
        # 检查模板中的占位符
        template_doc = Document(template_path)
        template_text = "\n".join([p.text for p in template_doc.paragraphs])
        template_placeholders = PlaceholderHandler.find_placeholders(template_text)
        
        # 显示占位符信息
        st.info(f"模板中包含以下占位符: {', '.join(template_placeholders)}")
        
        for index, row in df.iterrows():
            # 更新进度
            progress = (index + 1) / total_rows
            progress_bar.progress(progress)
            status_text.text(f"正在处理 {index+1}/{total_rows}...")
            
            # 准备替换数据
            replacements = row.to_dict()
            
            # 生成文件名
            output_filename = generate_output_filename(replacements, filename_template)
            output_path = os.path.join(st.session_state.output_dir, output_filename)
            
            # 处理文档
            if process_document(template_path, output_path, replacements):
                success_count += 1
                generated_files.append({
                    "name": output_filename,
                    "path": output_path
                })
        
        # 保存结果
        st.session_state.generated_files = generated_files
        st.session_state.processing_stage = 1
        st.session_state.filename_template = filename_template
        
        # 显示结果
        st.success(f"文档生成完成！成功: {success_count}/{total_rows}")
        if success_count < total_rows:
            st.warning(f"有 {total_rows - success_count} 个文档生成失败")
        
    except Exception as e:
        st.error(f"处理过程中发生错误: {str(e)}")
        st.session_state.processing_stage = 0

# 下载区域
if st.session_state.processing_stage == 1 and st.session_state.generated_files:
    st.header("3. 下载生成的文件")
    
    # 创建ZIP文件
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for file_info in st.session_state.generated_files:
            if os.path.exists(file_info["path"]):
                zip_file.write(file_info["path"], file_info["name"])
    
    zip_buffer.seek(0)
    
    # 提供下载按钮
    st.download_button(
        label=f"下载所有文档 (ZIP)",
        data=zip_buffer,
        file_name=f"generated_documents_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
        mime="application/zip"
    )
    
    # 显示生成的文件列表
    st.subheader("生成的文档列表")
    for file_info in st.session_state.generated_files:
        if os.path.exists(file_info["path"]):
            with open(file_info["path"], "rb") as f:
                st.download_button(
                    label=f"下载 {file_info['name']}",
                    data=f.read(),
                    file_name=file_info["name"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# 重置按钮
if st.button("重置系统"):
    # 清除session状态
    keys_to_clear = list(st.session_state.keys())
    for key in keys_to_clear:
        del st.session_state[key]
    
    # 重新初始化
    st.session_state.processing_stage = 0
    st.session_state.output_dir = ""
    st.session_state.generated_files = []
    st.session_state.filename_template = ""
    
    st.success("系统已重置，可以开始新的处理流程！")
    st.experimental_rerun()

# 使用说明
st.sidebar.header("使用说明")
st.sidebar.markdown("""
1. **上传文件**:
   - Excel数据文件（包含占位符数据）
   - Word模板文件（包含占位符标记）

2. **设置文件名模板**:
   - 使用与Word模板相同的占位符格式
   - 示例: `请款单（{{申请人}}-{{合计}}-集佳-{{日期}}）`

3. **开始生成**:
   - 点击"开始生成文档"按钮
   - 系统会处理所有数据行

4. **下载文档**:
   - 下载单个文档或所有文档(ZIP)
""")

st.sidebar.header("占位符格式")
st.sidebar.markdown("""
支持以下占位符格式:
- `{{key}}`
- `${key}`
- `{key}`
- `[[key]]`

在Word模板和文件名模板中使用相同的格式。
""")

st.sidebar.header("格式保留")
st.sidebar.markdown("""
系统会保留Word模板中的所有格式:
- 字体、大小、颜色
- 加粗、斜体、下划线
- 背景颜色
- 表格格式
""")
